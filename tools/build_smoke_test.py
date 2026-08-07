#!/usr/bin/env python3
"""Render every `docs/examples/<OWNER>/smoke_test.json` into its checklist.

    python tools/build_smoke_test.py             # markdown + docx
    python tools/build_smoke_test.py --no-docx   # markdown only

Outputs, next to the source:
    REVIT_SMOKE_TEST.md                          the reference copy
    <OWNER>_Revit_Smoke_Test_Checklist.docx      the printable session sheet

**These are outputs.** Edit the JSON. `tools/check_smoke_test.py` fails CI if the
committed markdown is not a byte-identical regeneration, which is what stops the
three-hand-maintained-copies problem coming back.

DOCX DEPENDENCY — the decision, stated
`python-docx` is present in this environment (1.2.0), so the Word document is
generated with it rather than by hand-rolling OOXML. The CI gate does NOT depend
on that: `check_smoke_test.py` asserts markdown freshness only, and this script
degrades to markdown-only with an explicit warning if python-docx is missing.
The alternative — emitting raw OOXML to avoid the dependency — buys nothing when
the gate never needs it, and costs a large amount of fragile XML.
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
import smoke_test_lib as L  # noqa: E402


# ── markdown ──────────────────────────────────────────────────────────────────

def _reach_note(s: dict) -> str:
    reach = s.get("reach")
    if reach == "button":
        return f"**{s['button']}** ({s['panel']} panel · {s['tab']} · {s['panelSection']})"
    if reach == "workflow":
        return f"run **{s['preset']}** (workflow only — no standalone button)"
    return "_Revit-native action — no STING command_"


def render_markdown(doc: dict) -> str:
    out: list[str] = []
    out.append(f"# {doc['title']}")
    out.append("")
    out.append("> **Generated file — do not edit.** Source:")
    out.append("> [`smoke_test.json`](smoke_test.json) · regenerate with")
    out.append("> `python tools/build_smoke_test.py` · gated by `tools/check_smoke_test.py`.")
    out.append("")
    if doc.get("intro"):
        out.append(doc["intro"])
        out.append("")

    steps = sorted(doc.get("steps") or [], key=lambda s: s["id"])
    section = None
    for s in steps:
        if s["section"] != section:
            section = s["section"]
            out.append(f"## {section}")
            out.append("")
        flags = []
        if s.get("preclearedOffline"):
            flags.append("pre-cleared offline")
        if s.get("optional"):
            flags.append("optional")
        flag_txt = f" _({', '.join(flags)})_" if flags else ""
        out.append(f"{s['id']}. **{s['title']}** — {_reach_note(s)}{flag_txt}")
        out.append(f"   - Expected: {s['expected']}")
        if s.get("commandTag"):
            out.append(f"   - Command tag: `{s['commandTag']}`")
        if s.get("fixture"):
            out.append(f"   - Fixture: `{s['fixture']}`")
        if s.get("artefact"):
            out.append(f"   - Artefact: `{s['artefact']}`")
        if s.get("dependsOn"):
            out.append(f"   - Depends on: step(s) {', '.join(str(d) for d in s['dependsOn'])}")
        if s.get("notes"):
            out.append(f"   - Note: {s['notes']}")
        out.append("")

    if doc.get("outro"):
        out.append(doc["outro"])
        out.append("")
    return "\n".join(out)


# ── docx ──────────────────────────────────────────────────────────────────────

def render_docx(doc: dict, path: Path) -> bool:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor
    except ImportError:
        print("  ! python-docx not installed — skipping the .docx "
              "(`pip install python-docx`). The markdown is still generated, and the "
              "CI gate only checks the markdown.")
        return False

    d = Document()
    for sect in d.sections:
        sect.left_margin = sect.right_margin = Pt(40)

    title = d.add_heading(doc["title"], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    warn = d.add_paragraph()
    r = warn.add_run("Generated from docs/examples/%s/smoke_test.json — do not edit this document; "
                     "edit the JSON and re-run tools/build_smoke_test.py." % doc["owner"])
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # ── Session header block ──
    d.add_heading("Session", level=1)
    hdr = d.add_table(rows=5, cols=2)
    hdr.style = "Table Grid"
    for i, label in enumerate(["Project / model", "Revit version + build", "Tester", "Date", "Plugin DLL path (from the live .addin)"]):
        hdr.cell(i, 0).text = label
        hdr.cell(i, 1).text = ""
        for p in hdr.cell(i, 0).paragraphs:
            for run in p.runs:
                run.bold = True

    if doc.get("intro"):
        d.add_heading("Before you start", level=1)
        for para in doc["intro"].split("\n\n"):
            if para.strip():
                d.add_paragraph(para.strip())

    # ── Steps ──
    steps = sorted(doc.get("steps") or [], key=lambda s: s["id"])
    section = None
    for s in steps:
        if s["section"] != section:
            section = s["section"]
            d.add_heading(section, level=1)

        t = d.add_table(rows=0, cols=2)
        t.style = "Table Grid"

        def row(k: str, v: str, bold_value: bool = False):
            cells = t.add_row().cells
            cells[0].text = k
            for p in cells[0].paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(9)
            cells[1].text = v
            for p in cells[1].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.bold = bold_value

        flags = []
        if s.get("preclearedOffline"):
            flags.append("pre-cleared offline")
        if s.get("optional"):
            flags.append("OPTIONAL")
        head = f"{s['id']}. {s['title']}"
        if flags:
            head += f"   [{', '.join(flags)}]"
        row("Step", head, bold_value=True)
        row("Run it via", _reach_note(s).replace("**", "").replace("_", ""))
        row("Expected", s["expected"])
        if s.get("commandTag"):
            row("Command tag", s["commandTag"])
        if s.get("fixture"):
            row("Fixture", s["fixture"])
        if s.get("artefact"):
            row("Artefact", s["artefact"])
        if s.get("dependsOn"):
            row("Depends on", "step(s) " + ", ".join(str(x) for x in s["dependsOn"]))
        if s.get("notes"):
            row("Note", s["notes"])
        # The tick box, as literal characters so no font substitution is needed.
        row("Result", "[  ] PASS      [  ] FAIL      [  ] NOT RUN")
        d.add_paragraph("")

    # ── Failure log ──
    d.add_heading("Failure log", level=1)
    d.add_paragraph("One row per failure. Attach the StingTools.log excerpt.")
    fl = d.add_table(rows=1, cols=4)
    fl.style = "Table Grid"
    for i, h in enumerate(["Step", "Command", "What happened", "StingTools.log excerpt / model context"]):
        fl.cell(0, i).text = h
        for p in fl.cell(0, i).paragraphs:
            for run in p.runs:
                run.bold = True
    for _ in range(8):
        fl.add_row()

    # ── Sign-off ──
    d.add_heading("Sign-off", level=1)
    so = d.add_table(rows=4, cols=4)
    so.style = "Table Grid"
    for i, h in enumerate(["Role", "Name", "Signature", "Date"]):
        so.cell(0, i).text = h
        for p in so.cell(0, i).paragraphs:
            for run in p.runs:
                run.bold = True
    for i, role in enumerate(["Tester", "BIM Manager", "Information Manager"], start=1):
        so.cell(i, 0).text = role

    if doc.get("outro"):
        d.add_heading("Notes", level=1)
        for para in doc["outro"].split("\n\n"):
            if para.strip():
                d.add_paragraph(para.strip())

    d.save(str(path))
    return True


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--repo-root", default=None)
    ap.add_argument("--no-docx", action="store_true", help="markdown only")
    args = ap.parse_args()

    root = Path(args.repo_root).resolve() if args.repo_root else Path(__file__).resolve().parent.parent
    sources = L.load_sources(root)
    if not sources:
        print("No docs/examples/*/smoke_test.json found — nothing to build.")
        return 1

    for owner, path, doc in sources:
        md_path = path.parent / "REVIT_SMOKE_TEST.md"
        md_path.write_text(render_markdown(doc), encoding="utf-8", newline="\n")
        print(f"  wrote {md_path.relative_to(root).as_posix()}")

        if not args.no_docx:
            docx_path = path.parent / f"{owner}_Revit_Smoke_Test_Checklist.docx"
            if render_docx(doc, docx_path):
                print(f"  wrote {docx_path.relative_to(root).as_posix()}")

    return 0


if __name__ == "__main__":
    sys.exit(main())
