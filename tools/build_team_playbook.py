# -*- coding: utf-8 -*-
"""Build the KUT Project Delivery Playbook as a formatted corporate .docx."""
import copy
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = 'KUT_Project_Delivery_Playbook.docx'

NAVY = RGBColor(0x1F, 0x36, 0x54)
SLATE = RGBColor(0x44, 0x4F, 0x5C)
GREY = RGBColor(0x7A, 0x7A, 0x7A)
RULE = '1F3654'
BAND = 'E8EDF3'
SHADE = 'F4F6F9'

d = Document()

# ── page + base styles ───────────────────────────────────────────────────────
sec = d.sections[0]
sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
sec.left_margin = sec.right_margin = Cm(2.2)
sec.top_margin = Cm(2.2)
sec.bottom_margin = Cm(2.0)

normal = d.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10)
normal.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12

for nm, size, colour, before, after in (
        ('Heading 1', 16, NAVY, 20, 8),
        ('Heading 2', 12.5, NAVY, 14, 5),
        ('Heading 3', 11, SLATE, 10, 4)):
    st = d.styles[nm]
    st.font.name = 'Calibri'
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = colour
    st.font.italic = False
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True


def _shade(cell, hexcolour):
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear')
    el.set(qn('w:fill'), hexcolour)
    cell._tc.get_or_add_tcPr().append(el)


def _rule_below(par, colour=RULE, size=8):
    pPr = par._p.get_or_add_pPr()
    bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), colour)
    bdr.append(bottom)
    pPr.append(bdr)


def h1(text, page_break=True):
    if page_break:
        d.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    p = d.add_heading(text, level=1)
    _rule_below(p)
    return p


def h2(text):
    return d.add_heading(text, level=2)


def h3(text):
    return d.add_heading(text, level=3)


def para(text, bold=False, italic=False, size=None, colour=None, align=None, space_after=None):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.bold, r.italic = bold, italic
    if size:
        r.font.size = Pt(size)
    if colour:
        r.font.color.rgb = colour
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(items, indent=0.0):
    for it in items:
        p = d.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6 + indent)
        p.paragraph_format.space_after = Pt(3)
        p.add_run(u'• ').bold = True
        _rich(p, it)


def numlist(items):
    for n, it in enumerate(items, 1):
        p = d.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.space_after = Pt(3)
        p.add_run('%d. ' % n).bold = True
        _rich(p, it)


def _rich(p, text):
    """**bold** segments inside a string."""
    for i, chunk in enumerate(text.split('**')):
        if chunk:
            p.add_run(chunk).bold = bool(i % 2)


def callout(text, title=None):
    t = d.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    c = t.cell(0, 0)
    _shade(c, SHADE)
    c._tc.get_or_add_tcPr().append(_margins())
    p0 = c.paragraphs[0]
    if title:
        r = p0.add_run(title.upper())
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = NAVY
        p0 = c.add_paragraph()
    _rich(p0, text)
    for r_ in p0.runs:
        r_.font.size = Pt(9.5)
    d.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def _margins():
    mar = OxmlElement('w:tcMar')
    for side, v in (('top', 100), ('start', 140), ('bottom', 100), ('end', 140)):
        e = OxmlElement('w:' + side)
        e.set(qn('w:w'), str(v))
        e.set(qn('w:type'), 'dxa')
        mar.append(e)
    return mar


def table(headers, rows, widths=None, font=8.5, caption=None):
    t = d.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0]
    for i, x in enumerate(headers):
        c = hdr.cells[i]
        c.text = ''
        p = c.paragraphs[0]
        r = p.add_run(x)
        r.bold = True
        r.font.size = Pt(font)
        r.font.color.rgb = NAVY
        _shade(c, BAND)
    trPr = hdr._tr.get_or_add_trPr()
    rep = OxmlElement('w:tblHeader')          # repeat header row across pages
    trPr.append(rep)
    for row in rows:
        cells = t.add_row().cells
        for i, x in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            _rich(p, str(x))
            for r in p.runs:
                r.font.size = Pt(font)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    if caption:
        cp = d.add_paragraph()
        cr = cp.add_run(caption)
        cr.font.size = Pt(8)
        cr.italic = True
        cr.font.color.rgb = GREY
    d.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def _field(par, instr):
    r = par.add_run()
    fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), 'begin'); r._r.append(fc)
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = instr
    r._r.append(it)
    fs = OxmlElement('w:fldChar'); fs.set(qn('w:fldCharType'), 'separate'); r._r.append(fs)
    fe = OxmlElement('w:fldChar'); fe.set(qn('w:fldCharType'), 'end'); r._r.append(fe)


# ── title page ───────────────────────────────────────────────────────────────
for _ in range(4):
    d.add_paragraph()
p = d.add_paragraph()
r = p.add_run('KAMPALA UGANDA TEMPLE')
r.font.size = Pt(11)
r.bold = True
r.font.color.rgb = SLATE
p.paragraph_format.space_after = Pt(2)

p = d.add_paragraph()
r = p.add_run('Project Delivery Playbook')
r.font.size = Pt(30)
r.bold = True
r.font.color.rgb = NAVY
p.paragraph_format.space_after = Pt(4)
_rule_below(p, RULE, 12)

p = d.add_paragraph()
r = p.add_run('Information management, production and delivery procedures for all appointed parties')
r.font.size = Pt(12)
r.font.color.rgb = SLATE
p.paragraph_format.space_before = Pt(10)

for _ in range(10):
    d.add_paragraph()

t = d.add_table(rows=0, cols=2)
for k, v in (('Document reference', 'KUT-PLN-ZZ-ZZ-RP-Z-0002'),
             ('Revision', 'P01'),
             ('Status / suitability', 'A1 — Authorised for use'),
             ('Prepared by', 'Planscape Consulting Engineers Ltd'),
             ('Role', 'Information Manager'),
             ('Date of issue', '[FILL]')):
    cells = t.add_row().cells
    rp = cells[0].paragraphs[0]
    rr = rp.add_run(k)
    rr.bold = True
    rr.font.size = Pt(9)
    rr.font.color.rgb = SLATE
    vp = cells[1].paragraphs[0]
    vr = vp.add_run(v)
    vr.font.size = Pt(9)
    cells[0].width, cells[1].width = Cm(5.0), Cm(11.6)

d.add_paragraph()
p = d.add_paragraph()
r = p.add_run('Prepared for the Kampala Uganda Temple project on behalf of the Lead Appointed Party. '
              'Issued through the Common Data Environment. Uncontrolled when printed.')
r.font.size = Pt(8)
r.italic = True
r.font.color.rgb = GREY

# ── footer ───────────────────────────────────────────────────────────────────
footer = sec.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
fr = fp.add_run('KUT Project Delivery Playbook   |   Rev P01   |   Page ')
fr.font.size = Pt(8)
fr.font.color.rgb = GREY
_field(fp, ' PAGE ')
fr2 = fp.add_run(' of ')
fr2.font.size = Pt(8)
fr2.font.color.rgb = GREY
_field(fp, ' NUMPAGES ')
for r in fp.runs:
    r.font.size = Pt(8)
    r.font.color.rgb = GREY

# ── document control ─────────────────────────────────────────────────────────
h1('Document control')

h2('Revision history')
table(['Rev', 'Date', 'Prepared', 'Checked', 'Summary of change'],
      [['P01', '[FILL]', '[FILL]', '[FILL]', 'First issue for mobilisation'],
       ['', '', '', '', ''],
       ['', '', '', '', '']],
      widths=[1.4, 2.4, 2.6, 2.6, 7.6])

h2('Distribution')
table(['Organisation', 'Role', 'Recipient'],
      [['The Church — Special Projects', 'Appointing Party', '[FILL]'],
       ['Symbion Consulting Group Studios', 'Lead Appointed Party', '[FILL]'],
       ['Planscape Consulting Engineers Ltd', 'Information Manager', 'Mayanja Davis'],
       ['[FILL]', 'Architecture / Interiors', '[FILL]'],
       ['[FILL]', 'Structural engineering', '[FILL]'],
       ['[FILL]', 'Mechanical, electrical and plumbing', '[FILL]'],
       ['[FILL]', 'Fire protection', '[FILL]'],
       ['[FILL]', 'Civil and site', '[FILL]'],
       ['[FILL]', 'Quantity surveying / cost', '[FILL]'],
       ['[FILL]', 'Contractor (from Stage 3.1)', '[FILL]']],
      widths=[6.4, 5.4, 4.8])

h2('Status of this document')
para('This playbook sets out how information is produced, checked, exchanged and accepted on the Kampala '
     'Uganda Temple project. It applies to every organisation and every individual producing project '
     'information, for the full duration of the appointment.')
para('The BIM Execution Plan is the contractual statement of the information requirements. This playbook is '
     'the working procedure that delivers them. Where the two documents differ, the BIM Execution Plan takes '
     'precedence and this playbook is corrected at the next revision.')
callout('Compliance with this playbook is a condition of information acceptance. Information that does not '
        'meet the requirements set out here will be returned unaccepted, and the originating party remains '
        'responsible for the programme consequences.', 'Compliance')

h2('Contents')
table(['Section', 'Title'],
      [['1', 'Purpose and how to use this playbook'],
       ['2', 'Project information'],
       ['3', 'Roles, responsibilities and authorities'],
       ['4', 'Information numbering'],
       ['5', 'The Common Data Environment'],
       ['6', 'Modelling standards'],
       ['7', 'Information requirements by stage'],
       ['8', 'Stage-by-stage delivery'],
       ['9', 'Operating rhythm and meetings'],
       ['10', 'Information delivery planning (MIDP and TIDP)'],
       ['11', 'Quality gates and acceptance'],
       ['12', 'Clash detection and coordination'],
       ['13', 'Specialist information streams'],
       ['14', 'Joining the project'],
       ['15', 'Change, risk and escalation'],
       ['Appendix A', 'Pre-share checklist'],
       ['Appendix B', 'Gate pack contents'],
       ['Appendix C', 'Kickoff agenda'],
       ['Appendix D', 'Project rules'],
       ['Appendix E', 'Classification code tables']],
      widths=[2.6, 14.0], font=9)

# ── 1 ────────────────────────────────────────────────────────────────────────
h1('1  Purpose and how to use this playbook')

h2('1.1  Purpose')
para('This playbook answers three questions for every party on the project: what information you are required '
     'to produce, how it must be structured and checked, and when it must be delivered. It converts the '
     'information requirements of the appointment into day-to-day working procedure.')

h2('1.2  Audience and reading route')
table(['If you are', 'Read in full', 'Keep to hand'],
      [['Newly appointed to the project', 'Sections 1 to 6, then your discipline in Section 8', 'Section 4 and Appendix A'],
       ['A Task Team Manager', 'The whole document', 'Sections 8, 9 and 11'],
       ['Producing models or drawings daily', 'Sections 4, 6 and 7', 'Appendix A'],
       ['The Information Manager', 'The whole document', 'Sections 9, 11 and 12'],
       ['The Contractor or a specialist subcontractor', 'Sections 1 to 5, 8.6 to 8.8, and 13', 'Appendices A and B']],
      widths=[5.0, 7.2, 4.4])

h2('1.3  The three governing rules')
para('Three rules underpin every procedure in this document. Where a procedure is unclear, these prevail.')
numlist([
    '**All information is issued through the Common Data Environment.** Information transmitted by email, '
    'messaging application or portable media has not been issued and carries no status.',
    '**No information is shared until it has passed the pre-share checklist** at Appendix A. The check is '
    'performed by the originating task team, not by the Information Manager.',
    '**The numbering system in Section 4 is mandatory and is not open to local interpretation.** A single '
    'incorrect container name propagates to the register, the transmittal, the coordination report and the '
    'handover data.',
])

# ── 2 ────────────────────────────────────────────────────────────────────────
h1('2  Project information')

h2('2.1  Project particulars')
table(['Item', 'Detail'],
      [['Project', 'Kampala Uganda Temple'],
       ['Project code', 'KUT'],
       ['Location', '[FILL — address, Kampala, Uganda]'],
       ['Appointing Party', 'The Church — Special Projects Department'],
       ['Lead Appointed Party', 'Symbion Consulting Group Studios'],
       ['Information Manager', 'Planscape Consulting Engineers Ltd'],
       ['Procurement route', '[FILL]'],
       ['Programme', '49 months. Phase 2 (design) 11 months; Phase 3 (construction and close-out) 38 months'],
       ['Key dates', '[FILL — appointment, Deliverables A to D, practical completion]']],
      widths=[5.0, 11.6])

h2('2.2  Project systems')
table(['Function', 'System', 'Authority'],
      [['Common Data Environment', 'Autodesk Construction Cloud', 'The single authoritative environment for all project information'],
       ['Model authoring', 'Autodesk Revit 2025', 'Millimetres; shared coordinate system fixed at mobilisation'],
       ['Coordination and clash', 'Navisworks and ACC Model Coordination', 'System of record for clash results; issues tracked as ACC Issues'],
       ['Design review', 'Bluebeam Studio', 'Owner review sessions; comment close-out is a gate condition'],
       ['FF&E, finishes and O&M', 'Fohlio', 'Owner single source of truth; the model links to it and does not duplicate it'],
       ['Specifications', 'RIB SpecLink, CSI MasterFormat', 'Specification authority; reconciled against the model at each gate'],
       ['Building management', 'Niagara (Tridium)', 'Operational system; reconciled to the model at handover'],
       ['Handover data', 'COBie 2.4', 'Structured asset data delivered with the record model']],
      widths=[3.8, 4.4, 8.4])

h2('2.3  Volumes')
para('The project is divided into seven volumes. The volume is the second most significant field in every '
     'container name and every asset identifier, and it governs how models are split and federated.')
table(['Volume code', 'Volume', 'Numbering value'],
      [['BLD1', 'Temple', '01'],
       ['BLD2', 'Meetinghouse', '02'],
       ['BLD3', 'Housing and ancillary', '03'],
       ['BLD4', 'Grounds', '04'],
       ['BLD5', 'Utility', '05'],
       ['BLD6', 'Guard house', '06'],
       ['EXT', 'Site-wide and external works', '00']],
      widths=[3.4, 7.6, 5.6])

h2('2.4  Stages and information gates')
table(['Stage', 'Name', 'Months', 'LOD', 'Gate'],
      [['0', 'Mobilisation', 'M0 to M1', '—', 'Kit issued, teams trained, CDE operational'],
       ['2.1', 'Basis of Design (Deliverable A)', 'M1', '200', 'Massing and generic systems coordinated'],
       ['2.2', 'Developed Design (Deliverable B, 50%)', 'M2 to M4', '300', 'Real geometry, correctly located'],
       ['2.3', 'Technical Design (Deliverable C, 100%)', 'M5 to M8', '350', 'Interfaces and connections resolved'],
       ['2.4', 'Tender', 'M9 to M10', '350', 'Tender set issued from the CDE'],
       ['2.5', 'Conformed set', 'M11', '350', 'Addenda incorporated and reissued'],
       ['3.1', 'Construction administration', 'M12 to M43', '400', 'Fabrication and installation-ready information'],
       ['3.2', 'FF&E installation', 'M40 to M43', '400', 'FF&E installed and reconciled'],
       ['3.3', 'Close-out (Deliverable D)', 'M44 to M45', '500', 'Verified record model and handover data']],
      widths=[1.6, 5.4, 2.6, 1.4, 5.6])
callout('Deliverable D is delivered at LOD 500, not LOD 400. LOD 500 requires verification that the modelled '
        'element corresponds to the element actually installed, together with its asset data. For serviceable '
        'plant this includes serial number and installation date, which must be captured progressively during '
        'Stage 3.1 and cannot be reconstructed at close-out. See Section 7.3.', 'Note')

# ── 3 ────────────────────────────────────────────────────────────────────────
h1('3  Roles, responsibilities and authorities')

h2('3.1  Roles')
table(['Role', 'Held by', 'Responsible for'],
      [['Appointing Party', 'The Church', 'Exchange information requirements; acceptance of each deliverable'],
       ['Lead Appointed Party', 'Symbion', 'The appointment; design leadership; chairing design and coordination meetings'],
       ['Information Manager', 'Planscape', 'The CDE, BEP, MIDP and standards; the quality gate; federation and clash management; registers, transmittals and handover data'],
       ['Task Team Manager', 'Each appointed party', 'The discipline model, the TIDP, data quality, and authorisation of every share'],
       ['Task team members', 'Each appointed party', 'Authoring to the standards in Sections 4, 6 and 7'],
       ['Quantity Surveyor', '[FILL]', 'Quantities and cost derived from the model'],
       ['Contractor', '[FILL]', 'As-built capture, commissioning information, specialist models'],
       ['Controls and commissioning contractor', '[FILL]', 'Niagara station, point naming, commissioning records'],
       ['Interior Designer', '[FILL]', 'FF&E and finishes design and the Fohlio record']],
      widths=[3.8, 3.4, 9.4])
callout('The Information Manager is accountable for information, not for design. Where coordination requires '
        'a design change, the responsible discipline makes that change and retains design responsibility. '
        'The Information Manager makes the issue visible, tracks it and confirms its closure.', 'Scope of the Information Manager')

h2('3.2  Responsibility matrix')
para('R — carries out the work.   A — accountable for the outcome.   C — consulted.   I — informed.', size=9, italic=True)
table(['Activity', 'IM', 'Lead AP', 'Arch', 'Struct', 'MEP', 'QS', 'Contr'],
      [['Maintain the Common Data Environment', 'A/R', 'A', 'I', 'I', 'I', 'I', 'C'],
       ['Issue and maintain the BIM Execution Plan', 'A/R', 'C', 'C', 'C', 'C', 'I', 'I'],
       ['Maintain the MIDP', 'A/R', 'C', 'C', 'C', 'C', 'C', 'C'],
       ['Produce and maintain a TIDP', 'C', 'A', 'R', 'R', 'R', 'R', 'R'],
       ['Author the discipline model', 'I', 'A', 'R', 'R', 'R', 'I', 'C'],
       ['Perform the pre-share check', 'C', 'I', 'R', 'R', 'R', 'I', 'R'],
       ['Federate the models', 'R', 'A', 'C', 'C', 'C', 'I', 'C'],
       ['Run clash detection', 'R', 'A', 'C', 'C', 'C', 'I', 'C'],
       ['Resolve a clash', 'C', 'A', 'R', 'R', 'R', 'I', 'R'],
       ['Chair the coordination meeting', 'C', 'A/R', 'C', 'C', 'C', 'I', 'C'],
       ['Produce drawings and sheets', 'C', 'A', 'R', 'R', 'R', 'I', 'C'],
       ['Drawing register and transmittals', 'A/R', 'C', 'C', 'C', 'C', 'I', 'I'],
       ['Quantities and bills of quantities', 'C', 'I', 'C', 'C', 'C', 'A/R', 'C'],
       ['FF&E and finishes data', 'R', 'A', 'R', 'I', 'I', 'C', 'C'],
       ['Specification reconciliation', 'R', 'A', 'R', 'R', 'R', 'C', 'I'],
       ['Gate audit and sign-off pack', 'A/R', 'A', 'C', 'C', 'C', 'C', 'C'],
       ['As-built capture', 'C', 'A', 'C', 'C', 'C', 'I', 'R'],
       ['Commissioning point list', 'R', 'A', 'I', 'I', 'C', 'I', 'R'],
       ['Handover data (COBie and O&M)', 'A/R', 'A', 'C', 'C', 'C', 'I', 'R']],
      widths=[7.0, 1.4, 1.6, 1.3, 1.4, 1.3, 1.2, 1.4], font=8)

# ── 4 ────────────────────────────────────────────────────────────────────────
h1('4  Information numbering')

para('Every model, drawing, document and modelled element on this project carries a structured identifier. '
     'The identifiers are machine-checked at every share and at every gate.')

h2('4.1  Container name')
para('Applies to every file, model, drawing, sheet, schedule and document.')
p = d.add_paragraph()
r = p.add_run('KUT - PLN - 01 - GF - M3 - A - 0001')
r.font.name = 'Consolas'
r.font.size = Pt(13)
r.bold = True
r.font.color.rgb = NAVY
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
table(['Field', 'Length', 'Meaning', 'Example'],
      [['Project', '3', 'Always KUT', 'KUT'],
       ['Originator', '3', 'The organisation that produced the container (Section 4.2)', 'PLN'],
       ['Volume', '2', 'The volume numbering value (Section 2.3). ZZ for all volumes', '01'],
       ['Level', '2', 'The level code (Section 4.5). ZZ for all levels, XX not applicable', 'GF'],
       ['Type', '2', 'The information type (Section 4.4)', 'M3'],
       ['Role', '1 to 2', 'The discipline (Section 4.3)', 'A'],
       ['Number', '4', 'Sequential within the set', '0001']],
      widths=[2.8, 1.8, 8.4, 3.6])
para('Fields are separated by a single hyphen. No spaces are permitted. All characters are upper case.')

callout('**Decision required before any information is numbered.** The automated check enforces an originator '
        'code of exactly three characters. Two options are available: issue three-character codes to every '
        'organisation, or extend the permitted range to three to six characters, which is closer to general '
        'ISO 19650 practice. The Appointing Party register determines which applies. No container may be '
        'numbered until this is confirmed in writing, because renumbering after Deliverable A affects every '
        'issued document.', 'Open item — Week 1')

h2('4.2  Originator codes')
table(['Organisation', 'Code'],
      [['Planscape Consulting Engineers Ltd', '[FILL]'],
       ['Symbion Consulting Group Studios', '[FILL]'],
       ['Architecture', '[FILL]'],
       ['Interiors', '[FILL]'],
       ['Structural engineering', '[FILL]'],
       ['Mechanical, electrical and plumbing', '[FILL]'],
       ['Fire protection', '[FILL]'],
       ['Civil and site', '[FILL]'],
       ['Quantity surveying', '[FILL]'],
       ['Contractor', '[FILL]'],
       ['Not known or not applicable', 'ZZZ']],
      widths=[11.6, 5.0])

h2('4.3  Role codes')
para('Only the codes below are valid on this project. A container or element carrying any other value will '
     'fail the standards audit.')
table(['Code', 'Discipline'],
      [['A', 'Architecture and interiors'], ['S', 'Structural'], ['M', 'Mechanical'],
       ['E', 'Electrical'], ['P', 'Plumbing and public health'], ['FP', 'Fire protection'],
       ['LV', 'Low voltage and communications'], ['G', 'Civil and site']],
      widths=[2.6, 14.0])

h2('4.4  Type codes')
table(['Code', 'Type', 'Code', 'Type'],
      [['M3', '3D model', 'RP', 'Report'],
       ['M2', '2D model or drafting', 'CA', 'Calculation'],
       ['DR', 'Drawing', 'RD', 'Room data sheet'],
       ['SH', 'Sheet', 'MS', 'Method statement'],
       ['SC', 'Schedule', 'PP', 'Presentation'],
       ['SP', 'Specification', 'CR', 'Clash or coordination report']],
      widths=[2.0, 6.3, 2.0, 6.3])

h2('4.5  Level codes')
table(['Code', 'Level'],
      [['B1', 'Basement level 1'], ['GF', 'Ground floor'],
       ['01, 02, 03 …', 'Upper floors'], ['RF', 'Roof'],
       ['ZZ', 'All levels, or not level-specific'], ['XX', 'Not applicable']],
      widths=[3.4, 13.2])
para('Level codes must correspond exactly to the level names in the issued project template. Local variants '
     'are not permitted.')

h2('4.6  Revision and suitability')
table(['Stage of the container', 'Revision series', 'Example'],
      [['Preliminary, before contract', 'P01, P02, P03 …', 'P02'],
       ['Contractual, published', 'C01, C02, C03 …', 'C01']],
      widths=[7.0, 5.6, 4.0])
para('The revision advances only when the container is re-issued through the CDE. Working saves do not consume '
     'a revision.')
table(['Suitability', 'Meaning', 'CDE state'],
      [['S0', 'Work in progress. Not for use by any other party', 'WIP'],
       ['S1', 'Shared for coordination', 'Shared'],
       ['S2', 'Shared for information', 'Shared'],
       ['S3', 'Shared for review and comment', 'Shared'],
       ['S4', 'Shared for stage approval', 'Shared'],
       ['A1 to An', 'Published and authorised. Contractual', 'Published'],
       ['B1 to Bn', 'Published and authorised with comments', 'Published']],
      widths=[3.0, 9.6, 4.0])
callout('A suitability code is a statement of the reliance other parties may place on the information. '
        'Applying a coordination or approval code to work that is not complete is a breach of this procedure '
        'and will be raised with the Lead Appointed Party.', 'Suitability')

h2('4.7  Asset identifier')
para('Every modelled element carries an eight-field identifier assembled from the data held on the element.')
p = d.add_paragraph()
r = p.add_run('M - BLD1 - Z01 - L02 - HVAC - SUP - AHU - 0003')
r.font.name = 'Consolas'
r.font.size = Pt(12)
r.bold = True
r.font.color.rgb = NAVY
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
table(['Field', 'Content', 'Source'],
      [['Discipline', 'Role code (Section 4.3)', 'Element category and discipline'],
       ['Location', 'Volume code (Section 2.3)', 'Room, workset or project information'],
       ['Zone', 'Zone code', 'Room or zone assignment'],
       ['Level', 'Level code (Section 4.5)', 'Element level'],
       ['System', 'System code (Appendix E)', 'The connected system'],
       ['Function', 'Function code (Appendix E)', 'The system role'],
       ['Product', 'Product code', 'Family and type'],
       ['Sequence', 'Four digits', 'Assigned in sequence within the volume']],
      widths=[3.0, 6.0, 7.6])

h3('What the task team must do')
para('The identifier is assembled from data already present in a correctly built model. It is not typed. '
     'Three conditions must be met for it to be correct.')
bullet([
    '**Every element must be attributable to a volume.** Adopt one of two conventions per model and confirm '
    'it with the Information Manager: per-volume worksets named BLD2_Mechanical, BLD3_Architecture and so on; '
    'or one model per volume with the volume set on project information.',
    '**Rooms must be placed and named before the first coordination share.** Room boundaries are the '
    'strongest source of location and zone data available.',
    '**MEP elements must be connected into real systems.** System data determines the system and function '
    'fields, the schedules and the commissioning point list.',
])
callout('An element that belongs to no room, no workset and no volume is assigned to BLD1 by default. It will '
        'be reported as a low-confidence assignment at every gate until it is corrected. The count of such '
        'elements is reported in the monthly status report.', 'Default assignment')

h2('4.8  Worked examples')
table(['Information', 'Container name'],
      [['Temple architectural model, all levels', 'KUT-XXX-01-ZZ-M3-A-0001'],
       ['Meetinghouse mechanical model', 'KUT-XXX-02-ZZ-M3-M-0001'],
       ['Temple ground floor general arrangement sheet', 'KUT-XXX-01-GF-SH-A-0100'],
       ['Site-wide drainage drawing', 'KUT-XXX-00-ZZ-DR-P-0050'],
       ['Federated coordination model', 'KUT-PLN-ZZ-ZZ-M3-Z-0001'],
       ['Coordination report, cycle 07', 'KUT-PLN-ZZ-ZZ-CR-Z-0007'],
       ['Temple level 2 room data sheet', 'KUT-XXX-01-02-RD-A-0210']],
      widths=[8.6, 8.0])

# ── 5 ────────────────────────────────────────────────────────────────────────
h1('5  The Common Data Environment')

h2('5.1  States')
table(['State', 'Purpose', 'Who may use the information', 'Suitability'],
      [['WIP', 'The originating task team\'s working area', 'The originating task team only', 'S0'],
       ['Shared', 'Coordination and review between parties', 'All appointed parties, for the purpose stated by the suitability code', 'S1 to S4'],
       ['Published', 'Authorised, contractual issue', 'All parties, for construction and contractual purposes', 'A1, B1'],
       ['Archived', 'Superseded information retained for the record', 'Reference only', 'As issued']],
      widths=[2.4, 4.6, 6.6, 3.0])
para('Information is never deleted. Superseded revisions are archived and remain retrievable.')

h2('5.2  Authorisation to change state')
table(['Transition', 'Authorised by', 'Preconditions'],
      [['WIP to Shared', 'Task Team Manager', 'Pre-share checklist completed and recorded (Appendix A)'],
       ['Shared to Published', 'Information Manager, on the authority of the Lead Appointed Party',
        'Gate audit passed; review comments closed; register updated'],
       ['Published to Archived', 'Information Manager', 'A superseding revision has been published']],
      widths=[3.6, 5.6, 7.4])

h2('5.3  Mandatory requirements')
bullet([
    '**All issue is through the CDE.** Email and messaging are for discussion only and confer no status on '
    'any information.',
    '**One shared coordinate system and project base point**, established at mobilisation and never altered. '
    'All models are linked by shared coordinates.',
    '**Units are millimetres.** Level and grid names are taken from the issued project template.',
    '**Content originates from the issued family library.** New content is submitted to the Information '
    'Manager for checking before use.',
    '**Imported CAD may be used as an underlay only.** It must not constitute deliverable geometry.',
    '**Security-minded information management applies (ISO 19650-5).** Project information, including images, '
    'plans and visualisations, must not be published, circulated outside the project, or used in promotional '
    'or portfolio material without the written permission of the Appointing Party.',
])

# ── 6 ────────────────────────────────────────────────────────────────────────
h1('6  Modelling standards')
table(['Subject', 'Requirement'],
      [['Origin and orientation', 'Shared coordinate system and true north from the issued template. The project base point and survey point must not be moved'],
       ['Units', 'Millimetres'],
       ['Levels and grids', 'As issued in the template. Any change requires the agreement of the Information Manager, as level codes appear in every container name'],
       ['Worksets', 'Per Section 4.7. Modelling on the default workset is not permitted'],
       ['Rooms', 'Placed, named and correctly bounded before the first coordination share'],
       ['Families', 'From the issued library. Loadable families must carry the project shared parameters'],
       ['Systems', 'MEP elements connected into real systems. System data drives identifiers, schedules and the commissioning point list'],
       ['Level of detail', 'Modelled to the stage requirement in Section 7. Over-modelling in early stages is to be avoided'],
       ['Phases', 'As issued. Local phases must not be created'],
       ['Linked models', 'Linked by shared coordinates, pinned, and never bound into the host model'],
       ['File hygiene', 'Purged before every share; file size reported in the share note'],
       ['Warnings', 'Reviewed before every share. No critical warnings at a gate']],
      widths=[4.0, 12.6])

# ── 7 ────────────────────────────────────────────────────────────────────────
h1('7  Information requirements by stage')

para('Level of Development defines how far an element has been developed and what data it carries. The '
     'requirements below are checked automatically at each gate. Geometry alone does not satisfy a stage; '
     'the data listed must be present on the element.')

h2('7.1  General requirements at each stage')
table(['LOD', 'Stage', 'Geometry', 'Data required on every element'],
      [['200', 'Deliverable A', 'Present. Generic or placeholder families permitted', 'Asset identifier'],
       ['300', 'Deliverable B', 'Present. Placeholder and generic families no longer permitted; a real type is required', 'Asset identifier'],
       ['350', 'Deliverable C and conformed set', 'As LOD 300', 'Asset identifier, product code'],
       ['400', 'Construction', 'As LOD 350. A manufacturer type is required', 'Asset identifier, product code, model reference'],
       ['500', 'Deliverable D', 'As LOD 400, verified against the installed element', 'As LOD 400, plus the asset data in Section 7.3']],
      widths=[1.4, 4.2, 6.0, 5.0])

h2('7.2  Additional requirements by category')
para('The following categories carry requirements beyond the general rule. Where a category is not listed, '
     'the general requirement in Section 7.1 applies.')
table(['Category', 'From LOD 300', 'From LOD 350', 'From LOD 400'],
      [['Mechanical equipment', 'System type', 'Product code', 'Model reference, manufacturer'],
       ['Electrical equipment', 'System type', 'Product code', 'Model reference, manufacturer'],
       ['Lighting fixtures', 'System type', 'Product code', 'Model reference, manufacturer'],
       ['Plumbing fixtures', 'System type', 'Product code', 'Model reference, manufacturer, **maintenance type**'],
       ['Air terminals, sprinklers, fire alarm devices', 'System type', 'Product code', 'Model reference, manufacturer'],
       ['Electrical fixtures', 'System type', 'Product code', 'Model reference, manufacturer'],
       ['Ducts, pipes, conduits, cable trays and their fittings', 'System type', '—', '—'],
       ['Doors and windows', '—', 'Product code', 'Model reference, manufacturer'],
       ['Casework, specialty equipment, furniture and furniture systems', '—', 'Product code', 'Model reference, manufacturer'],
       ['Curtain panels and mullions', '—', 'Product code', 'Model reference, manufacturer'],
       ['Walls, floors, roofs, stairs, structural framing, columns and foundations', '—', 'Product code', '—'],
       ['Ceilings, railings and ramps', '—', '—', 'Product code']],
      widths=[6.0, 3.4, 3.2, 4.0], font=8)
callout('Plumbing fixtures require a maintenance type value from LOD 400. This requirement did not apply at '
        'earlier stages and the construction gate is the first point at which it is tested. Plumbing task '
        'teams should confirm the value is being captured at the start of Stage 3.1.', 'Change at LOD 400')

h2('7.3  Asset data for handover (LOD 500)')
para('At Deliverable D the following categories carry additional asset data. This data is captured during '
     'construction and installation. It cannot be reconstructed at close-out.')
table(['Category', 'Additional data at LOD 500'],
      [['Mechanical equipment', 'Serial number, installation date'],
       ['Electrical equipment', 'Serial number, installation date'],
       ['Lighting fixtures', 'Serial number, installation date'],
       ['Plumbing fixtures', 'Serial number, installation date'],
       ['Air terminals', 'Serial number, installation date'],
       ['Sprinklers', 'Serial number, installation date'],
       ['Fire alarm devices', 'Serial number, installation date'],
       ['Specialty equipment', 'Serial number, installation date'],
       ['Furniture', 'Installation date, FF&E reference'],
       ['Furniture systems', 'Installation date, FF&E reference']],
      widths=[7.6, 9.0])
callout('Capture of serial numbers and installation dates begins at Stage 3.1 and is reported monthly from '
        'the first month of construction. A programme that leaves this to Stage 3.3 will not achieve '
        'Deliverable D within the 60-day period following furniture installation.', 'Programme requirement')

# ── 8 ────────────────────────────────────────────────────────────────────────
h1('8  Stage-by-stage delivery')

def stage(num, title, entry, activities, deliverables, exit_):
    h2(num + '  ' + title)
    para('Entry criteria: ' + entry, size=9, italic=True)
    table(['Party', 'Activities'], activities, widths=[4.2, 12.4])
    para('Deliverables', bold=True, size=9.5)
    para(deliverables)
    para('Gate / exit criteria', bold=True, size=9.5)
    para(exit_)


stage('8.1', 'Stage 0 — Mobilisation (M0 to M1)',
      'Appointment in place; the CDE provisioned.',
      [['Information Manager', 'Establish the CDE structure and permissions; fix the shared coordinate system, levels and grids; issue the project template, family library and title blocks; issue the BEP, this playbook and the MIDP; deliver the kickoff'],
       ['Lead Appointed Party', 'Confirm the design programme and the volume and level register; nominate Task Team Managers'],
       ['Each appointed party', 'Nominate a Task Team Manager; return a TIDP; confirm software versions; attend the kickoff; establish the WIP area'],
       ['Appointing Party', 'Issue exchange information requirements and the organisation standards; confirm the originator code register']],
      'BIM Execution Plan; MIDP; this playbook; project template, family library and title blocks; a TIDP from every '
      'appointed party; the CDE operational with permissions applied.',
      'Every appointed party has produced a test model from the issued template, shared it once through the CDE, '
      'and passed the pre-share checklist. An appointed party that has not completed this does not commence Stage 2.1.')

stage('8.2', 'Stage 2.1 — Basis of Design, Deliverable A (M1, LOD 200)',
      'Stage 0 gate passed.',
      [['Architecture', 'Massing, volumes, primary circulation and gross areas; rooms placed'],
       ['Structure', 'Primary grid, indicative frame and foundations'],
       ['MEP', 'Plant space allocation, primary distribution routes, indicative loads'],
       ['Civil and site', 'Site model, levels, access and drainage strategy'],
       ['Quantity Surveyor', 'First order cost estimate from the model'],
       ['Information Manager', 'First federation; first clash detection at gross level; baseline model health report; area and programme audit against the brief']],
      'Discipline models at LOD 200; federated model; area schedule reconciled to the brief; first coordination '
      'report; Deliverable A drawing set.',
      'Gross spatial clashes resolved; areas reconciled to the brief; naming and data audit passed; review comments closed.')

stage('8.3', 'Stage 2.2 — Developed Design, Deliverable B (M2 to M4, LOD 300)',
      'Deliverable A accepted.',
      [['Architecture and interiors', 'Real geometry correctly located; door and window schedules commenced; room finishes structured; FF&E commenced'],
       ['Structure', 'Sized members and foundations; penetrations coordinated with MEP'],
       ['MEP', 'Real equipment; sized primary distribution; plant rooms coordinated; risers fixed'],
       ['Fire and low voltage', 'Detection, suppression and containment strategies modelled'],
       ['Quantity Surveyor', 'Bill of quantities from the model'],
       ['Information Manager', 'Fortnightly coordination cycle in operation; drawing production; register and transmittal for the data drop; monthly status report; first FF&E and finishes exchange']],
      'Models at LOD 300; 50% drawing set; bill of quantities; coordination reports with issues closed; updated MIDP; '
      'Deliverable B transmittal.',
      'No unresolved high priority clashes; LOD 300 verification passed; naming and data audit passed; review comments '
      'closed; bill of quantities issued.')

stage('8.4', 'Stage 2.3 — Technical Design, Deliverable C (M5 to M8, LOD 350)',
      'Deliverable B accepted.',
      [['All disciplines', 'Interfaces and connections resolved; builders work and penetrations agreed; details modelled where they govern coordination'],
       ['Interiors', 'FF&E and finishes complete and reconciled with the Fohlio record; room data sheets for key spaces'],
       ['MEP', 'Systems complete and connected; equipment carrying manufacturer and model data; BMS points identified'],
       ['Specification lead', 'CSI sections assigned; SpecLink table of contents reconciled against the model'],
       ['Quantity Surveyor', 'Tender bill of quantities'],
       ['Information Manager', 'Full drawing production; LOD 350 verification; specification reconciliation; FF&E currency check; gate pack']],
      'Models at LOD 350; 100% drawing set; tender bill of quantities; room data sheets; FF&E schedule; specification '
      'reconciliation report; Deliverable C transmittal.',
      'LOD 350 verification passed; no unresolved clashes; specification gaps closed or formally accepted; FF&E linked; '
      'review comments closed.')

h2('8.5  Stages 2.4 and 2.5 — Tender and conformed set (M9 to M11)')
para('The tender set is issued from the CDE at suitability A1. Queries are raised and answered as formal requests '
     'for information and are logged. No model change is made except by instruction. Following award, addenda and '
     'tender stage changes are incorporated and the set is regenerated and reissued as the conformed baseline '
     'against which construction proceeds. The register must show every superseded revision as archived.')

stage('8.6', 'Stage 3.1 — Construction administration (M12 to M43, LOD 400)',
      'Conformed set published; contractor mobilised.',
      [['Contractor', 'Shop drawings and fabrication models; progressive as-built capture; requests for information through the CDE'],
       ['Design team', 'Responses to requests for information; revisions issued with revision data; site queries'],
       ['Specialist subcontractors', 'Fabrication level models linked into the federation'],
       ['Controls contractor', 'Niagara station build; point naming agreed against the model'],
       ['Information Manager', 'Monthly federation and clash detection; revision control; register maintenance; monthly status report; asset data reporting; preparation of the commissioning point list']],
      'Construction stage models at LOD 400; revision-controlled drawing issues; monthly coordination and status '
      'reports; progressive asset data capture.',
      'Construction information complete; as-built capture current to within one month; asset data capture on programme.')

h2('8.7  Stage 3.2 — FF&E installation (M40 to M43)')
para('FF&E is installed and reconciled item by item against the Fohlio record. Finishes are verified against the '
     'installed condition. The gate requires the FF&E schedule to be reconciled with no unlinked items, and O&M '
     'information to be collected.')

stage('8.8', 'Stage 3.3 — Close-out, Deliverable D (M44 to M45, LOD 500)',
      'Practical completion of the relevant works; within 60 days of furniture installation.',
      [['Contractor', 'Final as-built information; commissioning records; warranties and O&M documentation'],
       ['Controls contractor', 'Live Niagara station reconciled against the model equipment and points'],
       ['Design team', 'Verification that the record model reflects the constructed building'],
       ['Information Manager', 'LOD 500 verification; asset data completeness; COBie handover data; final register, transmittal and archive']],
      'Verified record model at LOD 500; asset and equipment register; reconciled BMS point register; COBie 2.4 and '
      'O&M documentation; final drawing register; project archive.',
      'LOD 500 verification passed; asset data complete; handover information accepted by the Appointing Party.')

# ── 9 ────────────────────────────────────────────────────────────────────────
h1('9  Operating rhythm and meetings')

h2('9.1  The delivery cycle')
table(['Frequency', 'Activity', 'Responsible', 'Output'],
      [['Daily', 'Authoring in WIP. No information leaves WIP without the pre-share check', 'Task teams', '—'],
       ['Weekly, Tuesday', 'Progress note posted to the CDE: what has changed, what is blocked, what is planned', 'Task Team Managers', 'Progress note'],
       ['Fortnightly, Wednesday', 'Share by 12:00; federation, clash detection and reporting by 17:00', 'All parties; Information Manager', 'Coordination report and issues'],
       ['Fortnightly, Friday', 'Coordination meeting', 'Lead Appointed Party (chair)', 'Minutes and issue assignments'],
       ['Monthly', 'Status report: model health, compliance, clash burn-down, review close-out, FF&E currency, asset data', 'Information Manager', 'Monthly report'],
       ['Per gate', 'Gate audit, sign-off pack, transmittal and publication', 'Information Manager', 'Gate pack'],
       ['Per data drop', 'MIDP updated and register reissued', 'Information Manager', 'MIDP and register']],
      widths=[2.8, 7.0, 3.4, 3.4])

h2('9.2  The fortnightly coordination cycle')
table(['When', 'Activity'],
      [['Monday and Tuesday', 'Authoring in WIP'],
       ['Wednesday 12:00', 'Share closes. Every party shares to the CDE at the appropriate suitability. Information not shared by this time is not included in the cycle'],
       ['Wednesday 12:00 to 17:00', 'Federation, clash detection, grouping and prioritisation'],
       ['Wednesday 17:00', 'Coordination report and issues issued'],
       ['Thursday and Friday', 'Parties review the issues assigned to them'],
       ['Friday', 'Coordination meeting'],
       ['Following week', 'Resolution, for inclusion in the next share']],
      widths=[4.6, 12.0])
callout('The coordination report is issued 48 hours before the meeting so that participants attend having read '
        'it. Meeting time is for resolving issues, not for discovering them.', 'The 48-hour rule')

h2('9.3  Meetings')
table(['Meeting', 'Frequency', 'Chair', 'Attendees', 'Purpose'],
      [['Coordination', 'Fortnightly', 'Lead AP', 'Task Team Managers, Information Manager', 'Resolve clashes and interfaces'],
       ['Design team', 'Weekly', 'Lead AP', 'Design leads', 'Design decisions'],
       ['Information management', 'Monthly', 'Information Manager', 'Task Team Managers', 'Standards, data quality, delivery planning'],
       ['Owner review', 'Per gate', 'Appointing Party', 'All', 'Review and comment'],
       ['Gate sign-off', 'Per gate', 'Lead AP', 'Appointing Party, Information Manager, discipline leads', 'Acceptance of the deliverable'],
       ['Site progress', 'Weekly from Stage 3.1', 'Contractor', 'Site team and design team', 'Construction issues']],
      widths=[3.0, 2.2, 2.4, 5.0, 4.0], font=8)

# ── 10 ───────────────────────────────────────────────────────────────────────
h1('10  Information delivery planning')

h2('10.1  Definitions')
para('The Task Information Delivery Plan (TIDP) is produced by each appointed party and lists the information '
     'that party will deliver, when, at what level of development, and in what format. It is owned by the Task '
     'Team Manager.')
para('The Master Information Delivery Plan (MIDP) aggregates every TIDP into the project master. It is owned by '
     'the Information Manager and is the single record of what is due, from whom, and when.')

h2('10.2  Production and maintenance')
table(['When', 'Action'],
      [['Mobilisation', 'Each appointed party returns a TIDP. The Information Manager aggregates these into the MIDP baseline'],
       ['At each stage commencement', 'TIDPs are reviewed and re-baselined for the stage'],
       ['At each data drop', 'Actual dates are recorded and status updated'],
       ['Monthly', 'The MIDP is reissued with the status report'],
       ['On change', 'Any change to a delivery date is agreed and the plan re-baselined. Dates are not adjusted without agreement']],
      widths=[4.6, 12.0])

h2('10.3  Required fields')
para('TIDPs are submitted using the issued template and the following fields.')
table(['Field', 'Field', 'Field'],
      [['Reference', 'Format', 'Actual date'],
       ['Discipline', 'Suitability', 'Responsible'],
       ['Originator', 'CDE state', 'TIDP reference'],
       ['Deliverable', 'Planned release month', 'Status'],
       ['Type', 'Planned date', 'Notes'],
       ['Stage and level of development', '', '']],
      widths=[5.6, 5.5, 5.5])

# ── 11 ───────────────────────────────────────────────────────────────────────
h1('11  Quality gates and acceptance')
para('A gate is passed when every condition below is satisfied and evidenced in the gate pack. The Information '
     'Manager conducts the audit; the Lead Appointed Party and the Appointing Party accept.')
table(['#', 'Condition', 'Evidence'],
      [['1', 'Every deliverable due for the stage is present at the required suitability', 'MIDP extract with actual dates'],
       ['2', 'Container naming and standards compliance', 'Standards audit report, no errors'],
       ['3', 'Data completeness for the stage', 'Completeness report with element counts'],
       ['4', 'Level of development verified for the stage', 'LOD verification report and schedule'],
       ['5', 'No unresolved high priority clashes', 'Coordination report with issue status'],
       ['6', 'Review comments closed', 'Review close-out report'],
       ['7', 'Model health within tolerance', 'Model health report'],
       ['8', 'Register and transmittal issued', 'Drawing register and transmittal receipt'],
       ['9', 'Specification reconciled (from Stage 2.3)', 'Specification gap report'],
       ['10', 'FF&E linked and current (from Stage 2.2)', 'FF&E currency report'],
       ['11', 'Asset data capture on programme (from Stage 3.1)', 'Asset data completeness report']],
      widths=[1.0, 9.0, 6.6])
callout('Every report submitted in a gate pack must state the number of elements examined. A report expressing '
        'a percentage without a population is not evidence of compliance and will not be accepted.', 'Evidence')

# ── 12 ───────────────────────────────────────────────────────────────────────
h1('12  Clash detection and coordination')

h2('12.1  Priorities')
table(['Priority', 'Definition', 'Resolution required'],
      [['P1 Critical', 'Hard clash between permanent elements, or a clash preventing construction sequence', 'Before the next gate, without exception'],
       ['P2 Major', 'Hard clash resolvable by rerouting or offset; access or maintenance space compromised', 'Within two cycles'],
       ['P3 Minor', 'Soft clash, tolerance or clearance issue', 'Before the stage gate'],
       ['P4 Observation', 'Recorded for information; no action required at present', 'Logged only']],
      widths=[2.6, 8.4, 5.6])

h2('12.2  Process')
numlist([
    'The Information Manager federates the shared models and runs clash detection.',
    'Clashes are grouped so that each issue represents one physical problem, not one geometric intersection, '
    'and are then prioritised.',
    'Each issue is assigned to a discipline with a required resolution date and is tracked as an ACC Issue.',
    'The coordination report is issued 48 hours before the coordination meeting.',
    'The meeting reviews open issues in the federated model. Every issue leaves the meeting with an owner and a date.',
    'Resolution is demonstrated in the next share and the issue is closed with evidence.',
])
para('An issue assigned to a discipline is a record of a coordination condition, not an assessment of '
     'performance. Issues identified late, or not raised at all, carry programme risk for every party.')

h2('12.3  Clearances and tolerances')
table(['Interface', 'Clearance'],
      [['Structure to MEP', '[FILL] mm hard clearance, plus maintenance access'],
       ['MEP to MEP', '[FILL] mm'],
       ['Maintenance access to plant', 'Per manufacturer requirement, minimum [FILL] mm'],
       ['Ceiling void services zoning', 'Per the agreed services zoning drawing']],
      widths=[7.0, 9.6])

# ── 13 ───────────────────────────────────────────────────────────────────────
h1('13  Specialist information streams')

h2('13.1  FF&E and finishes')
para('Fohlio is the Appointing Party\'s source of truth for FF&E, finishes and operation and maintenance '
     'information. The model carries a reference to the Fohlio record. It does not hold a duplicate of it.')
table(['Stage', 'Activity'],
      [['Mobilisation', 'Room finish and FF&E parameters bound; field mapping agreed; a single shared identifier agreed per element'],
       ['From Stage 2.2', 'Each cycle: finishes and FF&E exported in the agreed format; the Interior Designer enriches the record in Fohlio; enriched data returned to the model, matched by room number, with a difference report reviewed before any data is written'],
       ['Stage 2.3', 'Room data sheets and the FF&E schedule produced from the reconciled data'],
       ['Monthly', 'Currency check between the model and the Fohlio record, reported in the status report'],
       ['Stages 3.2 and 3.3', 'Installed FF&E reconciled; asset identifiers aligned for handover']],
      widths=[3.4, 13.2])
callout('The exchange is matched on room number. A room renumbered without notifying the Information Manager '
        'breaks the match for that room, and the break is silent. Room renumbering must be raised in advance.',
        'Room numbering')

h2('13.2  Specifications')
para('CSI MasterFormat sections are assigned to model elements from Stage 2.3. The SpecLink table of contents is '
     'reconciled against the model at each gate, producing three registers: specified but not modelled, modelled '
     'but not specified, and title mismatches. Each entry is closed or formally accepted before the gate is passed.')

h2('13.3  Building management system')
para('The model states the equipment and points that should exist. The Niagara station states what is installed '
     'and running. Alignment between the two constitutes the operational data set delivered to the Appointing Party.')
table(['Stage', 'Activity'],
      [['Stage 2.3', 'Serviceable MEP elements carry BMS data (point name, protocol, system). The point naming convention is agreed between the MEP team and the controls contractor'],
       ['Stage 3.1, approximately M40', 'The commissioning point list is produced from the model and issued to the controls contractor for loading'],
       ['Stage 3.3', 'The model equipment and points are reconciled against the live station. Differences are resolved with the controls contractor']],
      widths=[4.4, 12.2])

h2('13.4  Handover data')
para('COBie 2.4 is produced from the record model at close-out, covering facility, floor, space, type, component, '
     'system, spare, job and document data. Its quality is determined by the asset data captured during Stages 3.1 '
     'to 3.3, in accordance with Section 7.3.')

# ── 14 ───────────────────────────────────────────────────────────────────────
h1('14  Joining the project')
para('The following are completed before any modelling begins.')
numlist([
    'Obtain CDE access and confirm visibility of the WIP, Shared and Published areas.',
    'Read Sections 4, 5 and 6.',
    'Obtain the project template, family library and title blocks. Models are started from the issued template; '
    'existing project files are not migrated.',
    'Confirm the originator code and the volume and workset convention with the Task Team Manager.',
    'Produce a test model containing a correctly named and correctly located element, and share it once.',
    'Attend the next coordination meeting as an observer.',
])

# ── 15 ───────────────────────────────────────────────────────────────────────
h1('15  Change, risk and escalation')

h2('15.1  Changes to standards')
para('Any change to the numbering system, the project template, the family library or the level of development '
     'requirements is made through the Information Manager, recorded in the BIM Execution Plan, and issued to all '
     'parties. Local variation is not permitted. Where a requirement is impractical for a discipline, it is to be '
     'raised rather than worked around.')

h2('15.2  Escalation')
table(['Circumstance', 'Raise to', 'Timing'],
      [['A requirement is unclear or impractical', 'Information Manager', 'Immediately'],
       ['A share will be missed', 'Task Team Manager, then Information Manager', 'Before the deadline'],
       ['A clash cannot be resolved within the discipline', 'Coordination meeting', 'Within the same cycle'],
       ['A design decision is preventing information production', 'Lead Appointed Party', 'Within the same week'],
       ['A gate is at risk', 'Information Manager, then Lead Appointed Party and Appointing Party', 'Not less than two weeks before the gate']],
      widths=[6.6, 6.0, 4.0])

h2('15.3  Project risks under active management')
table(['Risk', 'Control'],
      [['Appointing Party standards issued after mobilisation, affecting naming or level of development',
        'Standards are held as project configuration rather than embedded in content, so adoption is a configuration change rather than rework'],
       ['Originator code length unresolved (Section 4.1)', 'To be confirmed in Week 1, before any container is numbered'],
       ['Weak volume attribution through absent rooms or worksets', 'Rooms placed before the first share; assignment confidence audited at every gate and reported monthly'],
       ['Divergence between the model and the FF&E record', 'Monthly currency check reported in the status report'],
       ['Asset data deferred to close-out', 'Asset data capture is a Stage 3.1 activity, reported monthly from the first month of construction'],
       ['As-built capture falling behind construction', 'Capture maintained to within one month, verified monthly'],
       ['Late delivery of specialist fabrication models', 'Named in the MIDP with dates from Stage 2.3 onward']],
      widths=[6.4, 10.2])

# ── appendices ───────────────────────────────────────────────────────────────
h1('Appendix A  Pre-share checklist')
para('Completed by the originating task team before any information is moved from WIP to Shared. Retained by the '
     'Task Team Manager and produced on request.')
table(['', 'Check'],
      [[u'☐', 'Model originates from the issued template; coordinate system unaltered'],
       [u'☐', 'Units in millimetres; levels and grids unmodified'],
       [u'☐', 'Correct workset and volume for every element'],
       [u'☐', 'Rooms placed and named'],
       [u'☐', 'MEP elements connected into real systems'],
       [u'☐', 'Content from the issued library; no imported CAD used as model geometry'],
       [u'☐', 'Model purged; file size recorded'],
       [u'☐', 'Revit warnings reviewed; no critical warnings'],
       [u'☐', 'Container named in accordance with Section 4'],
       [u'☐', 'Suitability code correctly applied (Section 4.6)'],
       [u'☐', 'Revision advanced'],
       [u'☐', 'Data completeness for the current stage verified (Section 7)'],
       [u'☐', 'Share note recorded: what has changed, and what remains unresolved']],
      widths=[1.2, 15.4])

h1('Appendix B  Gate pack contents', page_break=False)
bullet([
    'MIDP extract for the stage, with actual dates',
    'Standards and container naming audit report',
    'Data completeness report, stating the population examined',
    'Level of development verification report and schedule',
    'Coordination report with issue status and burn-down',
    'Model health report',
    'Review comment close-out report',
    'Drawing register and transmittal receipt',
    'Specification reconciliation report (from Stage 2.3)',
    'FF&E currency report (from Stage 2.2)',
    'Asset data completeness report (from Stage 3.1)',
])

h1('Appendix C  Kickoff agenda')
table(['Time', 'Item'],
      [['0:00', 'Project, programme, gates and project team'],
       ['0:20', 'Information management framework: the CDE, states, suitability and revisions'],
       ['0:35', 'Information numbering (Section 4), with worked examples and a practical exercise'],
       ['1:05', 'The CDE: structure, permissions, sharing and publication'],
       ['1:25', 'Break'],
       ['1:35', 'Modelling standards and the pre-share checklist'],
       ['2:00', 'The fortnightly cycle and the meeting programme'],
       ['2:20', 'Clash detection: priorities, grouping, assignment and closure'],
       ['2:40', 'Information delivery planning: what is required and when'],
       ['3:00', 'Discipline sessions'],
       ['4:00', 'Close']],
      widths=[2.4, 14.2])

h1('Appendix D  Project rules')
numlist([
    'Start from the issued template. Do not migrate an existing project file.',
    'Name every container in accordance with Section 4, first time.',
    'Place rooms before sharing.',
    'Connect MEP elements into real systems.',
    'Apply suitability codes accurately.',
    'Issue all information through the CDE.',
    'Complete the pre-share checklist before every share.',
    'Read the coordination report before the coordination meeting.',
    'Never move the project origin.',
    'Raise problems early.',
])

h1('Appendix E  Classification code tables')
para('The tables below list the codes applicable to this project. Codes outside these tables are not valid '
     'without the agreement of the Information Manager.')

h2('E.1  System codes')
table(['Code', 'System', 'Code', 'System'],
      [['HVAC', 'Heating, ventilation and air conditioning', 'LV', 'Low voltage electrical distribution'],
       ['DCW', 'Domestic cold water', 'LPS', 'Lightning protection'],
       ['DHW', 'Domestic hot water', 'COM', 'Communications'],
       ['HWS', 'Heating hot water', 'ICT', 'Information and communications technology'],
       ['SAN', 'Sanitary drainage', 'SEC', 'Security and access control'],
       ['RWD', 'Rainwater drainage', 'ARC', 'Architectural'],
       ['SWD', 'Surface water drainage', 'STR', 'Structural'],
       ['GAS', 'Gas distribution', 'IRR', 'Irrigation'],
       ['FP', 'Fire protection', 'RWH', 'Rainwater harvesting'],
       ['FLS', 'Fire and life safety detection', 'GEN', 'General or unclassified']],
      widths=[1.8, 6.5, 1.8, 6.5], font=8)

h2('E.2  Function codes')
table(['Code', 'Function', 'Code', 'Function'],
      [['SUP', 'Supply', 'PWR', 'Power distribution'],
       ['RTN', 'Return', 'DCW', 'Domestic cold water'],
       ['EXH', 'Extract or exhaust', 'DHW', 'Domestic hot water'],
       ['FRA', 'Fresh or outside air', 'SAN', 'Sanitary drainage'],
       ['HTG', 'Heating', 'RWD', 'Rainwater drainage'],
       ['FP', 'Fire protection', 'ARC', 'Architectural'],
       ['FLS', 'Fire and life safety', 'STR', 'Structural'],
       ['COM', 'Communications', 'GEN', 'General or unclassified']],
      widths=[1.8, 6.5, 1.8, 6.5], font=8)

h2('E.3  Zone codes')
para('Zone codes are assigned per volume in the form Z01, Z02 and so on, and are confirmed by the Information '
     'Manager at mobilisation. The zoning register is issued with the project template.')

d.add_paragraph()
p = d.add_paragraph()
r = p.add_run('End of document')
r.font.size = Pt(9)
r.italic = True
r.font.color.rgb = GREY
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── core properties ──────────────────────────────────────────────────────────
cp = d.core_properties
cp.title = 'KUT Project Delivery Playbook'
cp.subject = 'Kampala Uganda Temple — information management, production and delivery procedures'
cp.author = 'Planscape Consulting Engineers Ltd'
cp.last_modified_by = 'Planscape Consulting Engineers Ltd'
cp.category = 'Project procedure'
cp.comments = 'Rev P01. Issued through the Common Data Environment. Uncontrolled when printed.'

d.save(OUT)
print('saved:', OUT)
