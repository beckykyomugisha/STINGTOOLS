# -*- coding: utf-8 -*-
"""Shared house style for Planscape corporate documents.

One definition of the page setup, palette, headings, tables and callouts, so a
set of issued documents looks like one set. Both tools/build_team_playbook.py
and tools/build_bep.py build on this; a change here changes every document the
next time it is generated.

Usage:

    from corporate_docx import CorporateDoc
    doc = CorporateDoc()
    doc.title_page('Project Delivery Playbook', 'Kampala Uganda Temple', rows)
    doc.h1('1  Purpose')
    doc.para('...')
    doc.table(['A', 'B'], [['1', '2']])
    doc.save('Out.docx')
"""
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

NAVY = RGBColor(0x1F, 0x36, 0x54)
SLATE = RGBColor(0x44, 0x4F, 0x5C)
GREY = RGBColor(0x7A, 0x7A, 0x7A)
RULE = '1F3654'
BAND = 'E8EDF3'
SHADE = 'F4F6F9'
EM_SPACE = u' '   # list separator; matches the issued documents


class CorporateDoc:
    def __init__(self):
        self.d = Document()
        self._page()
        self._styles()

    # ── setup ────────────────────────────────────────────────────────────
    def _page(self):
        s = self.d.sections[0]
        s.page_width, s.page_height = Cm(21.0), Cm(29.7)
        s.left_margin = s.right_margin = Cm(2.2)
        s.top_margin = Cm(2.2)
        s.bottom_margin = Cm(2.0)
        self.section = s

    def _styles(self):
        n = self.d.styles['Normal']
        n.font.name = 'Calibri'
        n.font.size = Pt(10)
        n.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        n.paragraph_format.space_after = Pt(6)
        n.paragraph_format.line_spacing = 1.12
        for nm, size, colour, before, after in (
                ('Heading 1', 16, NAVY, 20, 8),
                ('Heading 2', 12.5, NAVY, 14, 5),
                ('Heading 3', 11, SLATE, 10, 4)):
            st = self.d.styles[nm]
            st.font.name = 'Calibri'
            st.font.size = Pt(size)
            st.font.bold = True
            st.font.italic = False
            st.font.color.rgb = colour
            st.paragraph_format.space_before = Pt(before)
            st.paragraph_format.space_after = Pt(after)
            st.paragraph_format.keep_with_next = True

    # ── primitives ───────────────────────────────────────────────────────
    @staticmethod
    def _shade(cell, colour):
        el = OxmlElement('w:shd')
        el.set(qn('w:val'), 'clear')
        el.set(qn('w:fill'), colour)
        cell._tc.get_or_add_tcPr().append(el)

    @staticmethod
    def _rule_below(par, colour=RULE, size=8):
        pPr = par._p.get_or_add_pPr()
        bdr = OxmlElement('w:pBdr')
        b = OxmlElement('w:bottom')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(size))
        b.set(qn('w:space'), '4')
        b.set(qn('w:color'), colour)
        bdr.append(b)
        pPr.append(bdr)

    @staticmethod
    def _cell_margins():
        mar = OxmlElement('w:tcMar')
        for side, v in (('top', 100), ('start', 140), ('bottom', 100), ('end', 140)):
            e = OxmlElement('w:' + side)
            e.set(qn('w:w'), str(v))
            e.set(qn('w:type'), 'dxa')
            mar.append(e)
        return mar

    @staticmethod
    def _rich(p, text):
        """Render **bold** segments within a string."""
        for i, chunk in enumerate(str(text).split('**')):
            if chunk:
                p.add_run(chunk).bold = bool(i % 2)

    @staticmethod
    def _field(par, instr):
        r = par.add_run()
        for kind, tag in (('begin', 'w:fldChar'), (None, 'w:instrText'), ('separate', 'w:fldChar'),
                          ('end', 'w:fldChar')):
            e = OxmlElement(tag)
            if kind:
                e.set(qn('w:fldCharType'), kind)
            else:
                e.set(qn('xml:space'), 'preserve')
                e.text = instr
            r._r.append(e)

    # ── content ──────────────────────────────────────────────────────────
    def h1(self, text, page_break=True):
        if page_break:
            self.d.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        p = self.d.add_heading(text, level=1)
        self._rule_below(p)
        return p

    def h2(self, text):
        return self.d.add_heading(text, level=2)

    def h3(self, text):
        return self.d.add_heading(text, level=3)

    def para(self, text, bold=False, italic=False, size=None, colour=None, align=None,
             space_after=None, rich=False):
        p = self.d.add_paragraph()
        if rich:
            self._rich(p, text)
        else:
            r = p.add_run(text)
            r.bold, r.italic = bold, italic
            if size:
                r.font.size = Pt(size)
            if colour:
                r.font.color.rgb = colour
        if rich and (size or italic or colour):
            for r in p.runs:
                if size:
                    r.font.size = Pt(size)
                if italic:
                    r.italic = True
                if colour:
                    r.font.color.rgb = colour
        if align is not None:
            p.alignment = align
        if space_after is not None:
            p.paragraph_format.space_after = Pt(space_after)
        return p

    def bullet(self, items, indent=0.0):
        for it in items:
            p = self.d.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6 + indent)
            p.paragraph_format.space_after = Pt(3)
            p.add_run(u'•' + EM_SPACE).bold = True
            self._rich(p, it)

    def numlist(self, items):
        for n, it in enumerate(items, 1):
            p = self.d.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_after = Pt(3)
            p.add_run('%d.' % n + EM_SPACE).bold = True
            self._rich(p, it)

    def callout(self, text, title=None):
        t = self.d.add_table(rows=1, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        c = t.cell(0, 0)
        self._shade(c, SHADE)
        c._tc.get_or_add_tcPr().append(self._cell_margins())
        p0 = c.paragraphs[0]
        if title:
            r = p0.add_run(title.upper())
            r.bold = True
            r.font.size = Pt(8.5)
            r.font.color.rgb = NAVY
            p0 = c.add_paragraph()
        self._rich(p0, text)
        for r in p0.runs:
            r.font.size = Pt(9.5)
        self.d.add_paragraph().paragraph_format.space_after = Pt(2)
        return t

    def table(self, headers, rows, widths=None, font=8.5, caption=None):
        t = self.d.add_table(rows=1, cols=len(headers))
        t.style = 'Table Grid'
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr = t.rows[0]
        for i, x in enumerate(headers):
            c = hdr.cells[i]
            c.text = ''
            r = c.paragraphs[0].add_run(x)
            r.bold = True
            r.font.size = Pt(font)
            r.font.color.rgb = NAVY
            self._shade(c, BAND)
        hdr._tr.get_or_add_trPr().append(OxmlElement('w:tblHeader'))
        for row in rows:
            cells = t.add_row().cells
            for i, x in enumerate(row):
                cells[i].text = ''
                p = cells[i].paragraphs[0]
                p.paragraph_format.space_after = Pt(2)
                self._rich(p, x)
                for r in p.runs:
                    r.font.size = Pt(font)
        if widths:
            for row in t.rows:
                for i, w in enumerate(widths):
                    row.cells[i].width = Cm(w)
        if caption:
            cp = self.d.add_paragraph()
            cr = cp.add_run(caption)
            cr.font.size = Pt(8)
            cr.italic = True
            cr.font.color.rgb = GREY
        self.d.add_paragraph().paragraph_format.space_after = Pt(2)
        return t

    def mono(self, text, size=13):
        p = self.d.add_paragraph()
        r = p.add_run(text)
        r.font.name = 'Consolas'
        r.font.size = Pt(size)
        r.bold = True
        r.font.color.rgb = NAVY
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p

    # ── document furniture ───────────────────────────────────────────────
    def title_page(self, title, eyebrow, control_rows, strapline, note):
        for _ in range(4):
            self.d.add_paragraph()
        p = self.d.add_paragraph()
        r = p.add_run(eyebrow.upper())
        r.font.size = Pt(11)
        r.bold = True
        r.font.color.rgb = SLATE
        p.paragraph_format.space_after = Pt(2)

        p = self.d.add_paragraph()
        r = p.add_run(title)
        r.font.size = Pt(30)
        r.bold = True
        r.font.color.rgb = NAVY
        p.paragraph_format.space_after = Pt(4)
        self._rule_below(p, RULE, 12)

        p = self.d.add_paragraph()
        r = p.add_run(strapline)
        r.font.size = Pt(12)
        r.font.color.rgb = SLATE
        p.paragraph_format.space_before = Pt(10)

        for _ in range(10):
            self.d.add_paragraph()

        t = self.d.add_table(rows=0, cols=2)
        for k, v in control_rows:
            cells = t.add_row().cells
            kr = cells[0].paragraphs[0].add_run(k)
            kr.bold = True
            kr.font.size = Pt(9)
            kr.font.color.rgb = SLATE
            vr = cells[1].paragraphs[0].add_run(v)
            vr.font.size = Pt(9)
            cells[0].width, cells[1].width = Cm(5.0), Cm(11.6)

        self.d.add_paragraph()
        p = self.d.add_paragraph()
        r = p.add_run(note)
        r.font.size = Pt(8)
        r.italic = True
        r.font.color.rgb = GREY

    def footer(self, text):
        fp = self.section.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fp.add_run(text + '   |   Page ')
        self._field(fp, ' PAGE ')
        fp.add_run(' of ')
        self._field(fp, ' NUMPAGES ')
        for r in fp.runs:
            r.font.size = Pt(8)
            r.font.color.rgb = GREY

    def end_mark(self, text='End of document'):
        self.d.add_paragraph()
        p = self.d.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(9)
        r.italic = True
        r.font.color.rgb = GREY
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def properties(self, title, subject, category, comments,
                   author='Planscape Consulting Engineers Ltd'):
        cp = self.d.core_properties
        cp.title = title
        cp.subject = subject
        cp.author = author
        cp.last_modified_by = author
        cp.category = category
        cp.comments = comments

    def save(self, path):
        self.d.save(path)
        return path
